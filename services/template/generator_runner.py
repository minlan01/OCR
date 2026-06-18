"""
生成器执行服务 — 安全执行用户上传的 Python 生成器代码，输出 Word 文档

执行逻辑:
  1. 如果模板有 generator_code → 在受限命名空间中执行，传入 data，获取生成文本
  2. 如果模板没有 generator_code → 用默认 docx_exporter 渲染提取后的数据
  3. 将生成文本转为格式化的 Word 文档（支持参考文书样式提取）
"""
from __future__ import annotations

import ast
import re
import tempfile
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from loguru import logger

_XML_INVALID_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

_DEFAULT_BODY_FONT = "仿宋"
_DEFAULT_BODY_SIZE = Pt(14)
_DEFAULT_BODY_ALIGNMENT = WD_ALIGN_PARAGRAPH.LEFT
_DEFAULT_TITLE_FONT = "黑体"
_DEFAULT_TITLE_SIZE = Pt(22)
_DEFAULT_HEADING_FONT = "黑体"
_DEFAULT_HEADING_SIZE = Pt(15)


def run_generator_to_docx(
    extracted_data: dict,
    generator_code: str | None,
    template_name: str = "",
    reference_doc: bytes | None = None,
) -> bytes:
    """执行生成器并返回 Word 文档字节

    Args:
        extracted_data: LLM 提取后的 JSON 数据
        generator_code: 用户上传的 Python 生成器代码
        template_name: 模板名称
        reference_doc: 参考文书 .docx 字节（用于提取样式）

    Returns:
        .docx 文件字节
    """
    if generator_code:
        generated_text = _execute_generator(generator_code, extracted_data)
    else:
        generated_text = _default_generate(extracted_data)

    styles = _DEFAULT_STYLES.copy()
    if reference_doc:
        ref_styles = _extract_styles_from_reference(reference_doc)
        styles.update(ref_styles)
        logger.info(f"Reference doc styles applied: font={styles['body_font']}, size={styles['body_size']}, align={styles['body_alignment']}")

    return _text_to_docx_bytes(generated_text, template_name, styles)


_DEFAULT_STYLES: dict[str, Any] = {
    "body_font": _DEFAULT_BODY_FONT,
    "body_size": _DEFAULT_BODY_SIZE,
    "body_alignment": _DEFAULT_BODY_ALIGNMENT,
    "title_font": _DEFAULT_TITLE_FONT,
    "title_size": _DEFAULT_TITLE_SIZE,
    "heading_font": _DEFAULT_HEADING_FONT,
    "heading_size": _DEFAULT_HEADING_SIZE,
    "left_margin": Cm(3.18),
    "right_margin": Cm(3.18),
    "top_margin": Cm(2.54),
    "bottom_margin": Cm(2.54),
}


def _extract_styles_from_reference(doc_bytes: bytes) -> dict[str, Any]:
    """从参考文书 .docx 中提取样式参数

    策略: 统计所有段落的字体/字号/对齐方式，取众数作为正文样式；
          标题样式取第一个居中段落的字体/字号。
    """
    styles: dict[str, Any] = {}
    try:
        doc = Document(Path(__file__).parent / "__ref_tmp__.docx" if False else None)
        import io
        doc = Document(io.BytesIO(doc_bytes))
    except Exception as e:
        logger.warning(f"Failed to open reference doc, using defaults: {e}")
        return styles

    font_counter: Counter[str] = Counter()
    size_counter: Counter[float] = Counter()
    align_counter: Counter[int] = Counter()
    title_font = None
    title_size = None

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        alignment = para.alignment
        if alignment is not None:
            align_counter[alignment] += 1

        for run in para.runs:
            if not run.text.strip():
                continue
            fn = run.font.name
            if fn:
                font_counter[fn] += 1
            sz = run.font.size
            if sz:
                size_counter[sz.pt] += 1

            if title_font is None and alignment == WD_ALIGN_PARAGRAPH.CENTER:
                title_font = fn
                title_size = sz

    if font_counter:
        styles["body_font"] = font_counter.most_common(1)[0][0]
    if size_counter:
        most_common_pt = size_counter.most_common(1)[0][0]
        styles["body_size"] = Pt(int(most_common_pt)) if most_common_pt == int(most_common_pt) else Pt(most_common_pt)
    if title_font:
        styles["title_font"] = title_font
    if title_size:
        styles["title_size"] = title_size

    for section in doc.sections:
        if section.left_margin:
            styles["left_margin"] = section.left_margin
        if section.right_margin:
            styles["right_margin"] = section.right_margin
        if section.top_margin:
            styles["top_margin"] = section.top_margin
        if section.bottom_margin:
            styles["bottom_margin"] = section.bottom_margin
        break

    return styles


# ── 安全：AST 白名单校验 ──
# 禁止 import、open、eval/exec、属性访问 __dunder__ 等

_FORBIDDEN_MODULES = {
    "os", "sys", "subprocess", "shutil", "pathlib",
    "socket", "http", "urllib", "requests",
    "ctypes", "multiprocessing", "threading",
    "pickle", "marshal", "importlib",
}
_FORBIDDEN_BUILTINS = {"__import__", "eval", "exec", "compile", "open", "globals", "locals"}

_MAX_CODE_SIZE = 256 * 1024  # 256KB


class _CodeSafetyValidator(ast.NodeVisitor):
    """遍历 AST 检查危险操作"""

    def __init__(self):
        self.violations: list[str] = []

    def visit_Import(self, node):
        for alias in node.names:
            root = alias.name.split(".")[0]
            if root in _FORBIDDEN_MODULES:
                self.violations.append(f"Import of '{alias.name}' is forbidden")
        self.generic_visit(node)

    def visit_ImportFrom(self, node):
        if node.module:
            root = node.module.split(".")[0]
            if root in _FORBIDDEN_MODULES:
                self.violations.append(f"Import from '{node.module}' is forbidden")
        self.generic_visit(node)

    def visit_Call(self, node):
        # 检查 __import__("os") 等变体
        if isinstance(node.func, ast.Name) and node.func.id in _FORBIDDEN_BUILTINS:
            self.violations.append(f"Call to '{node.func.id}()' is forbidden")
        self.generic_visit(node)

    def visit_Attribute(self, node):
        # 禁止 __dunder__ 属性访问（__class__, __subclasses__, __globals__ 等）
        if node.attr.startswith("__") and node.attr.endswith("__"):
            self.violations.append(f"Access to dunder attribute '{node.attr}' is forbidden")
        self.generic_visit(node)


def _validate_generator_code(code: str) -> None:
    """安全校验生成器代码：大小限制 + AST 白名单"""
    if len(code) > _MAX_CODE_SIZE:
        raise ValueError(
            f"Generator code too large: {len(code)} bytes (max {_MAX_CODE_SIZE // 1024}KB)"
        )
    try:
        tree = ast.parse(code)
    except SyntaxError as e:
        raise ValueError(f"Generator code has syntax error: {e}")

    validator = _CodeSafetyValidator()
    validator.visit(tree)
    if validator.violations:
        raise ValueError(
            "Generator code contains forbidden operations: " + "; ".join(validator.violations[:5])
        )


def _execute_generator(code: str, data: dict) -> str:
    """在受限命名空间中执行生成器代码

    约定:
      - 顶层 generate(data) -> str 函数，或
      - 类 __init__(self, data) + generate(self) -> str

    安全措施:
      - AST 白名单校验（禁止 import os/sys/subprocess 等）
      - 代码大小限制（256KB）
      - 禁止 __dunder__ 属性访问
      - 受限内置命名空间
    """
    # 安全校验
    _validate_generator_code(code)

    # 受限内置：移除危险函数
    safe_builtins = {k: v for k, v in __builtins__.__dict__.items() if k not in _FORBIDDEN_BUILTINS} if isinstance(__builtins__, dict) else \
        {k: getattr(__builtins__, k) for k in dir(__builtins__) if k not in _FORBIDDEN_BUILTINS and not k.startswith("_")}
    safe_builtins["__builtins__"] = safe_builtins  # 递归限制

    namespace: dict[str, Any] = {"__builtins__": safe_builtins}
    try:
        exec(code, namespace)
    except Exception as e:
        raise ValueError(f"Generator code compilation failed: {e}")

    generate_fn = namespace.get("generate")
    is_class_method = False

    if generate_fn is None:
        for cls_obj in namespace.values():
            if isinstance(cls_obj, type) and hasattr(cls_obj, "generate"):
                cls_instance = cls_obj(data)
                generate_fn = cls_instance.generate
                is_class_method = True
                break

    if generate_fn is None:
        raise ValueError(
            "Generator code must define a generate(data) function or a class with a generate() method"
        )

    try:
        if is_class_method:
            result = generate_fn()
        else:
            result = generate_fn(data)
    except Exception as e:
        raise ValueError(f"Generator execution failed: {e}")

    if not isinstance(result, str):
        raise ValueError(f"Generator must return str, got {type(result).__name__}")

    return result


def _default_generate(data: dict) -> str:
    """默认生成器：将提取数据格式化为纯文本"""
    import json
    return json.dumps(data, ensure_ascii=False, indent=2)


def _text_to_docx_bytes(text: str, title: str = "", styles: dict[str, Any] | None = None) -> bytes:
    """将文本内容转为格式化的 Word 文档"""
    if styles is None:
        styles = _DEFAULT_STYLES.copy()

    doc = Document()
    _setup_doc_styles(doc, styles)

    if title:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(_clean(title))
        run.font.size = styles.get("title_size", _DEFAULT_TITLE_SIZE)
        run.font.bold = True
        run.font.name = styles.get("title_font", _DEFAULT_TITLE_FONT)
        run.element.rPr.rFonts.set(qn("w:eastAsia"), styles.get("title_font", _DEFAULT_TITLE_FONT))
        para.space_after = Pt(16)

    for line in text.split("\n"):
        line = _clean(line)
        if not line:
            doc.add_paragraph()
            continue

        if _is_heading(line):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = styles.get("heading_size", _DEFAULT_HEADING_SIZE)
            run.font.bold = True
            run.font.name = styles.get("heading_font", _DEFAULT_HEADING_FONT)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), styles.get("heading_font", _DEFAULT_HEADING_FONT))
            para.space_before = Pt(12)
            para.space_after = Pt(6)
        else:
            para = doc.add_paragraph()
            para.alignment = styles.get("body_alignment", _DEFAULT_BODY_ALIGNMENT)
            para.paragraph_format.first_line_indent = Cm(0.74)
            run = para.add_run(line)
            run.font.size = styles.get("body_size", _DEFAULT_BODY_SIZE)
            run.font.name = styles.get("body_font", _DEFAULT_BODY_FONT)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), styles.get("body_font", _DEFAULT_BODY_FONT))

    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    import os
    os.close(fd)
    try:
        doc.save(tmp_path)
        return Path(tmp_path).read_bytes()
    finally:
        try:
            Path(tmp_path).unlink(missing_ok=True)
        except Exception:
            pass


def _setup_doc_styles(doc: Document, styles: dict[str, Any] | None = None):
    if styles is None:
        styles = _DEFAULT_STYLES.copy()

    body_font = styles.get("body_font", _DEFAULT_BODY_FONT)
    body_size = styles.get("body_size", _DEFAULT_BODY_SIZE)

    style = doc.styles["Normal"]
    font = style.font
    font.name = body_font
    font.size = body_size
    style.element.rPr.rFonts.set(qn("w:eastAsia"), body_font)

    pf = style.paragraph_format
    pf.alignment = styles.get("body_alignment", _DEFAULT_BODY_ALIGNMENT)

    for section in doc.sections:
        section.top_margin = styles.get("top_margin", Cm(2.54))
        section.bottom_margin = styles.get("bottom_margin", Cm(2.54))
        section.left_margin = styles.get("left_margin", Cm(3.18))
        section.right_margin = styles.get("right_margin", Cm(3.18))


def _is_heading(line: str) -> bool:
    """判断是否为标题行"""
    stripped = line.strip()
    if not stripped:
        return False
    heading_patterns = [
        r"^民事起诉状$",
        r"^诉讼请求$",
        r"^事实及理由$",
        r"^此致$",
        r"^第[一二三四五六七八九十]+[章节条款]",
        r"^[一二三四五六七八九十]+[、.]",
        r"^附表",
    ]
    for pattern in heading_patterns:
        if re.match(pattern, stripped):
            return True
    return False


def _clean(text: str) -> str:
    return _XML_INVALID_RE.sub("", text).strip()
