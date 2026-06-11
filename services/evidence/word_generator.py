"""
Word 文档生成器 — 使用 docxtpl 模板填充
生成立案证据、民事起诉状、司法鉴定申请书等 Word 文档

v2: 支持多原告、5段结构化事实与理由、赔偿总额留空
"""
from __future__ import annotations

import io
import uuid
from pathlib import Path
from typing import Any

from loguru import logger

from config.settings import settings
from services.evidence.classifier import CATEGORY_NAMES
from services.utils.date_utils import normalize_date

# 模板目录
TEMPLATE_DIR = Path(__file__).resolve().parent.parent.parent / "templates" / "evidence"
MINIO_BUCKET = "scan-result"


def _get_catalog_and_analysis(case_id: str) -> tuple[dict, dict]:
    """获取案件的清单数据和分析结果（线程安全，使用 run_in_worker）"""
    import uuid

    from db.models_evidence import EvidenceCase
    from db.session import get_session_factory, run_in_worker

    async def _fetch():
        from sqlalchemy import select

        case_uuid = uuid.UUID(case_id)
        async with get_session_factory()() as db:
            stmt = select(EvidenceCase).where(EvidenceCase.id == case_uuid)
            result = await db.execute(stmt)
            case = result.scalar_one_or_none()
            if not case:
                raise ValueError(f"Case not found: {case_id}")
            return case.catalog_data or {}, case.analysis_result or {}

    return run_in_worker(_fetch())


def _upload_to_minio(case_id: str, doc_bytes: bytes, filename: str, content_type: str) -> str:
    """上传到 MinIO 并返回 object_key"""
    from services.storage.minio_client import minio_client

    minio_key = f"evidence/{case_id}/{uuid.uuid4()}_{filename}"
    minio_client.upload_bytes(
        bucket=MINIO_BUCKET,
        object_key=minio_key,
        data=doc_bytes,
        content_type=content_type,
    )
    return minio_key


# ═══════════════════════════════════════════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════════════════════════════════════════

def _build_legacy_plaintiff(context: dict) -> dict:
    """从旧的单原告字段构建 plaintiff dict（向后兼容）"""
    return {
        "name": context.get("原告姓名1", context.get("plaintiff_name", "")),
        "relationship": context.get("亲属关系1", ""),
        "is_patient": context.get("is_patient1", False),
        "gender": context.get("性别1", ""),
        "ethnicity": context.get("民族1", ""),
        "birth_date": context.get("出生年月日1", ""),
        "address": context.get("住址1", ""),
        "id_number": context.get("身份证号1", ""),
        "phone": context.get("律师电话1", ""),
    }


def _normalize_ethnicity(value: str) -> str:
    """规范化民族字段：'汉' → '汉族'，'壮' → '壮族'"""
    if not value:
        return ""
    value = value.strip()
    # 常见民族简称
    short_names = {"汉": "汉族", "壮": "壮族", "回": "回族", "满": "满族",
                   "蒙古": "蒙古族", "藏": "藏族", "苗": "苗族", "彝": "彝族",
                   "土家": "土家族", "布依": "布依族", "侗": "侗族", "瑶": "瑶族",
                   "白": "白族", "朝鲜": "朝鲜族", "哈尼": "哈尼族", "黎": "黎族"}
    if value in short_names:
        return short_names[value]
    if not value.endswith("族"):
        return value + "族"
    return value


def _normalize_birth_date(value: str) -> str:
    """规范化出生日期字段：'1958-07-13' → '1958年7月13日'

    使用统一的 date_utils.normalize_date() 处理所有格式
    """
    if not value:
        return ""
    return normalize_date(value.strip())


def _extract_identity_from_catalog(catalog_data: dict, plaintiff_name: str) -> dict:
    """从 catalog_data 的身份证材料中提取完整身份信息（备选方案）

    当 LLM 没有正确提取身份信息时，直接从 OCR 文本中用正则提取。
    返回: {id_number, ethnicity, birth_date, address, gender}
    """
    import re

    result = {}

    # 正则模式 - 更宽松以匹配各种OCR输出格式
    id_pattern = re.compile(r'[1-9]\d{5}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]')
    ethnicity_pattern = re.compile(r'民\s*族[：:\s]*([\u4e00-\u9fff]{1,8})')
    gender_pattern = re.compile(r'性\s*别[：:\s]*(男|女)')
    birth_pattern_cn = re.compile(r'(?:出\s*生|出生日期|生\s*日)[：:\s]*((?:19|20)\d{2})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日')
    birth_pattern_iso = re.compile(r'(?:出\s*生|出生日期|生\s*日)[：:\s]*((?:19|20)\d{2})[-/.](\d{1,2})[-/.](\d{1,2})')
    birth_pattern_bare = re.compile(r'((?:19|20)\d{2})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日')
    address_pattern = re.compile(r'(?:住\s*址|住\s*所|地\s*址)[：:\s]*([\u4e00-\u9fff\d号楼栋室层区路街道巷村镇省市区县]{5,80})')

    def _matches_name(item: dict, name: str) -> bool:
        """多路径姓名匹配：identity.patient_name / original_filename"""
        if not name:
            return False
        # 路径1: identity.patient_name
        identity = item.get("identity", {})
        item_name = identity.get("patient_name", "")
        if name in item_name:
            return True
        # 路径2: evidence_name.original_filename（文件名通常含人名）
        ev_name = item.get("evidence_name", {})
        filename = ev_name.get("original_filename", "") or item.get("title", "")
        # 去掉扩展名和"正面"/"反面"等后缀
        clean_fn = re.sub(r'\.(jpg|jpeg|png|pdf|jpx)$', '', filename or "", flags=re.IGNORECASE)
        clean_fn = re.sub(r'(正|反)面$', '', clean_fn)
        if name in clean_fn:
            return True
        # 路径3: raw_extracted.layer_2_identity.patient_name
        raw = item.get("raw_extracted", {})
        layer2 = raw.get("layer_2_identity", {}) or {}
        raw_name = layer2.get("patient_name", "")
        if name in raw_name:
            return True
        return False

    def _fill_from_identity(identity: dict) -> None:
        """从 identity 层补充已有数据"""
        if not result.get("id_number"):
            pid = identity.get("patient_id", "")
            if pid and len(pid) >= 15 and "*" not in pid:
                result["id_number"] = pid

        if not result.get("birth_date"):
            pb = identity.get("patient_birth", "")
            if pb:
                result["birth_date"] = _normalize_birth_date(pb)

        if not result.get("address"):
            pa = identity.get("patient_address", "")
            if pa:
                result["address"] = pa

        if not result.get("gender"):
            pg = identity.get("patient_gender", "")
            if pg:
                result["gender"] = pg

        if not result.get("ethnicity"):
            pe = identity.get("ethnicity", "")
            if pe:
                result["ethnicity"] = _normalize_ethnicity(pe)

    def _fill_from_raw_extracted(raw: dict) -> None:
        """从 raw_extracted.layer_2_identity 补充（原始未脱敏数据）"""
        layer2 = raw.get("layer_2_identity", {}) or {}
        if not layer2:
            return
        if not result.get("id_number"):
            pid = layer2.get("patient_id", "")
            if pid and len(pid) >= 15 and "*" not in pid:
                result["id_number"] = pid

        if not result.get("birth_date"):
            pb = layer2.get("patient_birth", "")
            if pb:
                result["birth_date"] = _normalize_birth_date(pb)

        if not result.get("address"):
            pa = layer2.get("patient_address", "")
            if pa:
                result["address"] = pa

        if not result.get("gender"):
            pg = layer2.get("patient_gender", "")
            if pg:
                result["gender"] = pg

    def _fill_from_ocr_text(ocr_text: str) -> None:
        """从 OCR 文本中正则提取"""
        if not ocr_text:
            return

        # 提取身份证号
        if not result.get("id_number"):
            match = id_pattern.search(ocr_text)
            if match:
                result["id_number"] = match.group(0)

        # 提取性别
        if not result.get("gender"):
            match = gender_pattern.search(ocr_text)
            if match:
                result["gender"] = match.group(1)

        # 提取民族
        if not result.get("ethnicity"):
            match = ethnicity_pattern.search(ocr_text)
            if match:
                result["ethnicity"] = _normalize_ethnicity(match.group(1).strip())

        # 提取出生日期（支持多种格式）
        if not result.get("birth_date"):
            match = birth_pattern_cn.search(ocr_text)
            if match:
                result["birth_date"] = f"{match.group(1)}年{match.group(2)}月{match.group(3)}日"
            else:
                match = birth_pattern_iso.search(ocr_text)
                if match:
                    result["birth_date"] = f"{match.group(1)}年{int(match.group(2))}月{int(match.group(3))}日"
                else:
                    match = birth_pattern_bare.search(ocr_text)
                    if match:
                        result["birth_date"] = f"{match.group(1)}年{match.group(2)}月{match.group(3)}日"

        # 提取住址
        if not result.get("address"):
            match = address_pattern.search(ocr_text)
            if match:
                result["address"] = match.group(1).strip()

    for group in catalog_data.get("groups", []):
        if group.get("category") != "identity_id_card":
            continue

        for item in group.get("items", []):
            # 多路径姓名匹配
            if not _matches_name(item, plaintiff_name):
                continue

            # 依次从三个来源补充：identity 层 → raw_extracted → OCR 正则
            identity = item.get("identity", {})
            _fill_from_identity(identity)

            raw = item.get("raw_extracted", {})
            _fill_from_raw_extracted(raw)

            ocr_text = item.get("ocr_text", "")
            _fill_from_ocr_text(ocr_text)

    return result


def _get_plaintiffs(context: dict, catalog_data: dict = None) -> list[dict]:
    """从 context 中获取原告列表（支持新旧格式）"""
    plaintiffs = context.get("plaintiffs", [])
    if plaintiffs:
        # 先规范化已有数据格式
        for p in plaintiffs:
            if p.get("ethnicity"):
                p["ethnicity"] = _normalize_ethnicity(p["ethnicity"])
            if p.get("birth_date"):
                p["birth_date"] = _normalize_birth_date(p["birth_date"])

        # 备选方案：如果 LLM 没有提取完整身份信息，从 catalog_data 的身份证OCR中补充
        if catalog_data:
            for p in plaintiffs:
                name = p.get("name", "")
                if not name:
                    continue
                # 检查是否有缺失字段
                missing_fields = []
                for field in ("id_number", "ethnicity", "birth_date", "address", "gender"):
                    if not p.get(field):
                        missing_fields.append(field)
                if missing_fields:
                    catalog_identity = _extract_identity_from_catalog(catalog_data, name)
                    for field in missing_fields:
                        value = catalog_identity.get(field, "")
                        if value:
                            p[field] = value
        return plaintiffs

    # 旧格式：检测是否有多个原告
    plaintiffs = []
    i = 1
    while True:
        name_key = f"原告姓名{i}"
        if i == 1:
            name = context.get(name_key, context.get("plaintiff_name", ""))
        else:
            name = context.get(name_key, "")
        if not name and i > 1:
            break
        if not name:
            break
        plaintiffs.append({
            "name": name,
            "relationship": context.get(f"亲属关系{i}", ""),
            "is_patient": context.get(f"is_patient{i}", False),
            "gender": context.get(f"性别{i}", ""),
            "ethnicity": context.get(f"民族{i}", ""),
            "birth_date": context.get(f"出生年月日{i}", ""),
            "address": context.get(f"住址{i}", ""),
            "id_number": context.get(f"身份证号{i}", ""),
            "phone": context.get(f"律师电话{i}", context.get("律师电话1", "")),
        })
        i += 1

    if not plaintiffs:
        plaintiffs = [_build_legacy_plaintiff(context)]

    # 先规范化已有数据格式
    for p in plaintiffs:
        if p.get("ethnicity"):
            p["ethnicity"] = _normalize_ethnicity(p["ethnicity"])
        if p.get("birth_date"):
            p["birth_date"] = _normalize_birth_date(p["birth_date"])

    # 备选方案：如果 LLM 没有提取完整身份信息，从 catalog_data 的身份证OCR中补充
    if catalog_data:
        for p in plaintiffs:
            name = p.get("name", "")
            if not name:
                continue
            missing_fields = []
            for field in ("id_number", "ethnicity", "birth_date", "address", "gender"):
                if not p.get(field):
                    missing_fields.append(field)
            if missing_fields:
                catalog_identity = _extract_identity_from_catalog(catalog_data, name)
                for field in missing_fields:
                    value = catalog_identity.get(field, "")
                    if value:
                        p[field] = value

    return plaintiffs


# ═══════════════════════════════════════════════════════════════════════════════
# 立案证据
# ═══════════════════════════════════════════════════════════════════════════════

def generate_filing_evidence_inline_data(catalog_data: dict, analysis_result: dict) -> bytes | None:
    """纯数据驱动的立案证据生成 → 返回 bytes（不访问 DB / MinIO）"""
    return _build_filing_evidence_docx(catalog_data, analysis_result)


def _build_filing_evidence_docx(catalog_data: dict, analysis_result: dict) -> bytes:
    """构建立案证据 DOCX（纯数据驱动）"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 构建 context
    context: dict[str, Any] = {
        "case_name": analysis_result.get("case_name", ""),
        "case_type_name": analysis_result.get("case_type_name", ""),
        "plaintiff_name": analysis_result.get("原告姓名1", analysis_result.get("plaintiff_name", "")),
        "defendant_name": analysis_result.get("被告医院全称", analysis_result.get("defendant_name", "")),
        "catalog_items": [],
        "fee_items": [],
        "fee_summary": analysis_result.get("fee_summary", {}),
    }

    for group in catalog_data.get("groups", []):
        for item in group.get("items", []):
            context["catalog_items"].append({
                "index": item.get("index", 0),
                "title": item.get("title", ""),
                "proof_purpose": item.get("proof_purpose", ""),
            })

    fee_summary = catalog_data.get("fee_summary", {})
    for name, amount in fee_summary.items():
        try:
            amount = float(amount)
        except (ValueError, TypeError):
            continue
        context["fee_items"].append({"name": name, "amount": f"{amount:.2f}"})

    total_amt = catalog_data.get('total_amount', 0)
    try:
        total_amt = float(total_amt)
    except (ValueError, TypeError):
        total_amt = 0.0
    context["fee_summary"] = {"total_amount": f"{total_amt:.2f}"}

    # 尝试 docxtpl 模板
    template_path = TEMPLATE_DIR / "filing_evidence.docx"
    if template_path.exists():
        from docxtpl import DocxTemplate
        doc = DocxTemplate(str(template_path))
        doc.render(context)
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    # 无模板时用 python-docx 内联生成
    doc = Document()
    title = doc.add_heading("立案证据", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"案件名称：{context.get('case_name', '')}")
    doc.add_paragraph(f"案件类型：{context.get('case_type_name', '')}")
    doc.add_paragraph(f"原告：{context.get('plaintiff_name', '')}")
    doc.add_paragraph(f"被告：{context.get('defendant_name', '')}")

    doc.add_heading("证据材料清单", level=1)
    for item in context.get("catalog_items", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{item['index']}. {item['title']}")
        run.bold = True
        doc.add_paragraph(f"   证明目的：{item.get('proof_purpose', '')}")

    if context.get("fee_items"):
        doc.add_heading("赔偿费用汇总", level=1)
        for fee in context["fee_items"]:
            doc.add_paragraph(f"{fee['name']}：{fee['amount']}元")
        total = context.get("fee_summary", {}).get("total_amount", "0.00")
        p = doc.add_paragraph()
        run = p.add_run(f"合计：{total}元")
        run.bold = True

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def generate_filing_evidence(case_id: str) -> str:
    """生成立案证据 Word 文档 → 返回 MinIO key"""
    catalog_data, analysis_result = _get_catalog_and_analysis(case_id)
    doc_bytes = _build_filing_evidence_docx(catalog_data, analysis_result)

    minio_key = _upload_to_minio(
        case_id, doc_bytes, "立案证据.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    logger.info(f"Filing evidence generated: case={case_id} key={minio_key}")
    return minio_key


# ═══════════════════════════════════════════════════════════════════════════════
# 民事起诉状（v2: 多原告 + 5段结构化 + 赔偿总额留空）
# ═══════════════════════════════════════════════════════════════════════════════

def generate_complaint(case_id: str) -> str:
    """生成民事起诉状 Word 文档 → 返回 MinIO key"""
    catalog_data, analysis_result = _get_catalog_and_analysis(case_id)
    doc_bytes = _build_complaint_docx(catalog_data, analysis_result)

    minio_key = _upload_to_minio(
        case_id, doc_bytes, "民事起诉状.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    logger.info(f"Complaint generated: case={case_id} key={minio_key}")
    return minio_key


def generate_complaint_inline_data(catalog_data: dict, analysis_result: dict, lawyer_info: list[dict[str, str]] | None = None) -> bytes | None:
    """纯数据驱动的民事起诉状生成 → 返回 bytes"""
    return _build_complaint_docx(catalog_data, analysis_result, lawyer_info=lawyer_info or [])


def _year_to_chinese(year_str: str) -> str:
    """将阿拉伯数字年份转为中文大写，如 2026 → 二〇二六"""
    digit_map = {'0': '〇', '1': '一', '2': '二', '3': '三', '4': '四',
                 '5': '五', '6': '六', '7': '七', '8': '八', '9': '九'}
    return ''.join(digit_map.get(ch, ch) for ch in str(year_str))


import re

# LLM生成失败的占位文本模式
ERROR_PLACEHOLDER_PATTERN = re.compile(r'\[(段落生成失败|生成失败)[^\]]*\]')
PENDING_PLACEHOLDER_PATTERN = re.compile(r'\[待补充\]|（待补充）|\(待补充\)')


def _has_error_placeholder(text: str) -> bool:
    """检查文本是否包含LLM生成失败的占位标记"""
    if not text:
        return False
    return bool(ERROR_PLACEHOLDER_PATTERN.search(text))


def _has_pending_placeholder(text: str) -> bool:
    """检查文本是否包含待补充的占位标记"""
    if not text:
        return False
    return bool(PENDING_PLACEHOLDER_PATTERN.search(text))


def validate_analysis_result(analysis_result: dict) -> list[str]:
    """验证分析结果中是否有生成失败的段落

    Returns:
        list[str]: 失败的段落key列表，空列表表示全部正常
    """
    failed_sections = []
    paragraph_keys = ["paragraph_1", "paragraph_2", "paragraph_3", "paragraph_4", "paragraph_5"]

    for key in paragraph_keys:
        text = analysis_result.get(key, "")
        if _has_error_placeholder(text):
            failed_sections.append(key)

    return failed_sections


def validate_pending_fields(analysis_result: dict) -> list[str]:
    """验证分析结果中是否有待补充的字段

    Returns:
        list[str]: 待补充的字段描述列表，空列表表示全部正常
    """
    pending_fields = []

    # 检查关键字段
    key_fields = [
        ("defendant_name", "被告医院名称"),
        ("plaintiff_name", "原告姓名"),
        ("court_name", "受理法院"),
    ]

    for field_key, field_name in key_fields:
        value = analysis_result.get(field_key, "")
        if isinstance(value, str) and _has_pending_placeholder(value):
            pending_fields.append(field_name)

    # 检查段落中是否有待补充内容
    paragraph_keys = ["paragraph_1", "paragraph_2", "paragraph_3", "paragraph_4", "paragraph_5"]
    for key in paragraph_keys:
        text = analysis_result.get(key, "")
        if isinstance(text, str) and _has_pending_placeholder(text):
            pending_fields.append(f"{key} 段落")

    return pending_fields


def validate_required_fields(analysis_result: dict) -> list[str]:
    """验证文档生成所需的必要字段

    Returns:
        list[str]: 缺失的字段描述列表，空列表表示全部正常
    """
    missing = []

    # 检查原告信息
    plaintiffs = analysis_result.get("plaintiffs", [])
    if not plaintiffs:
        # 检查旧格式
        if not analysis_result.get("plaintiff_name") and not analysis_result.get("原告姓名1"):
            missing.append("原告姓名")

    # 检查被告信息
    defendant_name = (
        analysis_result.get("defendant_name")
        or analysis_result.get("被告医院全称")
        or ""
    )
    if not defendant_name:
        missing.append("被告医院名称")

    # 法院信息：不作为硬性阻断，缺失时在文档中留空让用户手填
    # （court_name 来自 LLM 提取，证据材料中可能不包含法院信息）

    return missing


def _clean_paragraph_text(text: str) -> str:
    """清理段落文本：移除段号标题、资质判断语句、猜测性语句等不应出现在正文中的内容"""
    if not text:
        return text

    # 移除段号标题（如 "四、司法鉴定意见\n\n"、"一、入院情况\n" 等）
    text = re.sub(r'^[一二三四五六七八九十]+[、．.][^\n]*\n*', '', text)
    
    # 移除加粗引导句（如 "**原告...诉被告...一案，现就...陈述如下：**\n"）
    text = re.sub(r'\*\*[^*]+\*\*\s*\n*', '', text)
    
    # 移除 "该机构具备法定鉴定资质。" 等资质判断语句
    text = re.sub(r'该机构具备法定鉴定资质[。，]?\s*', '', text)
    
    # 移除 "关于因果关系判定...未予明确界定...需补充鉴定" 等猜测性语句（多种变体）
    text = re.sub(r'关于因果关系判定[、，]医疗过错认定及其参与度[，,]该鉴定意见未予明确界定[，,]故尚需结合其他证据或补充鉴定予以确认[。，]?\s*', '', text)
    text = re.sub(r'但该鉴定意见未就[^。]*?(因果关系|过错|参与度)[^。]*?作出[^。]*?认定[。，]?\s*', '', text)
    text = re.sub(r'鉴于[^。]*?尚待[^。]*?明确[，,]?\s*恳请[^。]*?[。，]?\s*', '', text)
    text = re.sub(r'关于因果关系判定[、，][^。]*?(未予明确界定|未就[^。]*?作出明确认定)[^。]*?[。，]?\s*', '', text)
    
    # 移除 "本案在诉前调解阶段，经原告与被告XXX共同委托，" 等诉讼程序性描述
    text = re.sub(r'本案在诉前调解阶段[，,]经[^，,]+?[与和][^，,]+?共同委托[，,]\s*', '', text)
    text = re.sub(r'经原告[^，,]*?[与和]被告[^，,]*?共同委托[，,]\s*', '', text)
    
    # 移除报告编号中的方括号年份（如 [2026]）
    text = re.sub(r'\[\d{4}\]', '', text)

    # 开头的 "死者" 替换为 "患者"
    text = re.sub(r'^死者', '患者', text)
    
    # 清理多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    
    return text


def _build_complaint_docx(catalog_data: dict, analysis_result: dict, lawyer_info: list[dict[str, str]] | None = None) -> bytes:
    """构建民事起诉状 DOCX（多原告 + 5段结构化 + 无总额行 + 无证据清单）"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    context: dict[str, Any] = dict(analysis_result)
    if "evidence_list" not in context and "catalog_items" in context:
        context["evidence_list"] = context["catalog_items"]

    # 格式常量
    BODY_FONT = "仿宋"
    BODY_SIZE = Pt(14)  # 四号
    TITLE_FONT = "黑体"
    TITLE_SIZE = Pt(22)  # 二号
    LEFT_INDENT = Cm(0.05)
    RIGHT_INDENT = Cm(0.2)
    FIRST_INDENT = Cm(0.99)

    def _set_run_font(run, font_name: str, size):
        """设置 run 的字体（含东亚字体）和字号"""
        run.font.name = font_name
        run.font.size = size
        rPr = run._element.get_or_add_rPr()
        rFonts_elem = rPr.find(qn('w:rFonts'))
        if rFonts_elem is None:
            rFonts_elem = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts_elem)
        rFonts_elem.set(qn('w:eastAsia'), font_name)

    def _set_body_para(p, indent=True):
        """设置正文段落格式"""
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = LEFT_INDENT
        p.paragraph_format.right_indent = RIGHT_INDENT
        if indent:
            p.paragraph_format.first_line_indent = FIRST_INDENT

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 默认字体：仿宋四号
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    # 设置默认样式的东亚字体
    style_rPr = style.element.find(qn('w:rPr'))
    if style_rPr is None:
        style_rPr = OxmlElement('w:rPr')
        style.element.append(style_rPr)
    style_rFonts = style_rPr.find(qn('w:rFonts'))
    if style_rFonts is None:
        style_rFonts = OxmlElement('w:rFonts')
        style_rPr.insert(0, style_rFonts)
    style_rFonts.set(qn('w:eastAsia'), BODY_FONT)

    # === 标题 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("民事起诉状")
    _set_run_font(run, TITLE_FONT, TITLE_SIZE)
    run.font.bold = True

    # === 原告信息（支持多个）===
    # 未成年人案件：患者本人用"原告："，父母用"法定代理人："
    # 成年人案件：全部用"原告："
    is_minor = context.get("is_minor", False)

    plaintiffs = _get_plaintiffs(context, catalog_data)

    # ── 未成年人案件：修复多人 is_patient=true 的误标 ──
    # 如果多个原告都被标为 is_patient=true/is_patient 缺失且 relationship="本人"，
    # 只有最年轻的才是真正的患者，其他人需要推断关系
    if is_minor:
        patient_count = sum(1 for p in plaintiffs if p.get("is_patient", False) or p.get("relationship") == "本人")
        if patient_count > 1:
            import re as _re_local
            def _birth_year(p):
                m = _re_local.search(r'(\d{4})', p.get("birth_date", ""))
                return int(m.group(1)) if m else 0
            sorted_ps = sorted(plaintiffs, key=_birth_year, reverse=True)
            youngest = sorted_ps[0]
            patient_surname = (youngest.get("name") or "")[0] if youngest.get("name") else ""
            for p in plaintiffs:
                if p is youngest:
                    p["is_patient"] = True
                    p["relationship"] = "本人"
                else:
                    p["is_patient"] = False
                    rel = p.get("relationship", "")
                    if not rel or rel == "本人":
                        gender = p.get("gender", "")
                        p_surname = (p.get("name") or "")[0]
                        if p_surname == patient_surname:
                            p["relationship"] = "母亲" if "女" in gender else "父亲"
                        else:
                            p["relationship"] = "父亲" if "男" in gender else "母亲"

    for i, p_info in enumerate(plaintiffs):
        name = p_info.get("name") or ""
        rel = p_info.get("relationship") or ""
        gender = p_info.get("gender") or ""
        ethnicity = p_info.get("ethnicity") or ""
        birth = p_info.get("birth_date") or ""
        address = p_info.get("address") or ""
        id_number = p_info.get("id_number") or ""
        phone = p_info.get("phone") or ""

        is_patient = p_info.get("is_patient", False) or rel == "本人"

        # ── 判断是否为法定代理人（基于关系关键词） ──
        # 即使 is_patient 被误标，只要关系是父母/监护人，就应该是法定代理人
        _guardian_keywords = {"父亲", "母亲", "父", "母", "监护人", "爷爷", "奶奶", "外公", "外婆"}
        is_guardian = is_minor and rel and (rel in _guardian_keywords or any(kw in rel for kw in _guardian_keywords))

        # ── 标签选择：成年人全用"原告"，未成年人区分 ──
        if is_guardian or (is_minor and not is_patient and rel and rel != "本人"):
            # 未成年人的父母/法定监护人 → "法定代理人"
            label = "法定代理人："
        else:
            label = "原告："

        # ── 关系括号 ──
        # 参考文件格式："王文雄（系原告父亲）" — 用"系原告XX"而非"系患者XX"
        # 法定代理人一定有关系描述，即使 is_patient 被误标
        if rel and rel != "本人" and (not is_patient or is_guardian):
            rel_str = f"（系原告{rel}）"
        else:
            rel_str = ""

        parts = [f"{label}{name}{rel_str}"]
        if gender:
            parts.append(gender)
        if ethnicity:
            parts.append(ethnicity)
        if birth:
            parts.append(f"{birth}生")

        text = "，".join(parts)

        if address:
            text += f"，住址：{address}"
        if id_number:
            text += f"，公民身份号码：{id_number}"

        # 联系电话：优先使用律师电话，无律师时才用个人电话
        # 参考文件格式："联系电话：18487301173（凡律师）、15877881929（刘律师）"
        lawyer_parts = []
        if lawyer_info:
            for lw in lawyer_info:
                lw_name = lw.get("name", "").strip()
                lw_phone = lw.get("phone", "").strip()
                if lw_name and lw_phone:
                    lawyer_parts.append(f"{lw_phone}（{lw_name}律师）")
                elif lw_phone:
                    lawyer_parts.append(lw_phone)
        if lawyer_parts:
            text += "，联系电话：" + "、".join(lawyer_parts)
        elif phone:
            text += f"，联系电话：{phone}"

        para = doc.add_paragraph()
        _set_body_para(para)
        run_label = para.add_run(label)
        _set_run_font(run_label, BODY_FONT, BODY_SIZE)
        run_label.font.bold = True
        run_content = para.add_run(text[len(label):])
        _set_run_font(run_content, BODY_FONT, BODY_SIZE)

    # === 被告信息 ===
    defendant_name = context.get("defendant_name", context.get("被告医院全称", "")) or ""
    legal_rep = context.get("legal_representative", context.get("法定代表人姓名", "")) or ""
    credit_code = context.get("credit_code", context.get("统一社会信用代码", "")) or ""
    def_address = context.get("defendant_address", context.get("医院地址", "")) or ""
    def_phone = context.get("defendant_phone", context.get("医院电话", "")) or ""

    # 被告名称（独立行）
    para = doc.add_paragraph()
    _set_body_para(para)
    run_label = para.add_run("被告：")
    _set_run_font(run_label, BODY_FONT, BODY_SIZE)
    run_label.font.bold = True
    run_content = para.add_run(defendant_name)
    _set_run_font(run_content, BODY_FONT, BODY_SIZE)

    # 法定代表人（独立行）
    if legal_rep:
        para2 = doc.add_paragraph()
        _set_body_para(para2)
        run2 = para2.add_run(f"法定代表人：{legal_rep}")
        _set_run_font(run2, BODY_FONT, BODY_SIZE)

    # 统一社会信用代码
    if credit_code:
        para3 = doc.add_paragraph()
        _set_body_para(para3)
        run3 = para3.add_run(f"统一社会信用代码：{credit_code}")
        _set_run_font(run3, BODY_FONT, BODY_SIZE)

    # 地址（始终显示）
    para4 = doc.add_paragraph()
    _set_body_para(para4)
    run4 = para4.add_run(f"地址：{def_address}" if def_address else "地址：")
    _set_run_font(run4, BODY_FONT, BODY_SIZE)

    # 联系电话（始终显示，即使为空）
    para5 = doc.add_paragraph()
    _set_body_para(para5)
    run5 = para5.add_run(f"联系电话：{def_phone}" if def_phone else "联系电话：")
    _set_run_font(run5, BODY_FONT, BODY_SIZE)

    # === 诉讼请求（固定模板，金额挖空）===
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_body_para(heading, indent=False)
    run = heading.add_run("诉讼请求")
    _set_run_font(run, BODY_FONT, BODY_SIZE)
    run.font.bold = True

    # 第一条：赔偿请求（根据案件类型选择赔偿项目，金额从赔偿计算结果填充）
    case_type = context.get("case_type", "death")
    total_amount = context.get("compensation_total_amount", 0)
    try:
        total_amount = float(total_amount)
    except (ValueError, TypeError):
        total_amount = 0.0

    # 格式化金额为中文大写 + 阿拉伯数字
    if total_amount > 0:
        amount_str = f"{total_amount:,.2f}"
        amount_display = f"{amount_str}元"
    else:
        amount_display = "______元"

    if case_type == "injury" and is_minor:
        # 未成年人/新生儿案件：无误工费，含鉴定费、交通住宿费、残疾赔偿金
        claim_items = f"医疗费、护理费、住院伙食补助费、营养费、残疾赔偿金、交通住宿费、鉴定费、精神损害抚慰金等暂计人民币{amount_display}"
        retention_clause = "因本案尚未鉴定，上述费用待鉴定后再行补充变更"
    elif case_type == "injury":
        # 成年人伤残案件
        claim_items = f"医疗费、误工费、护理费、住院伙食补助费、营养费、伤残赔偿金（包含被扶养人生活费）、后续治疗费、交通费、精神损害抚慰金等暂计人民币{amount_display}"
        retention_clause = "因本案尚未鉴定，上述赔偿费用待鉴定后再行补充变更"
    elif case_type == "neonatal":
        # 旧 neonatal 兼容（已合并为 injury + is_minor，此处兜底）
        claim_items = f"医疗费、护理费、住院伙食补助费、营养费、残疾赔偿金、交通住宿费、鉴定费、精神损害抚慰金等暂计人民币{amount_display}"
        retention_clause = "因本案尚未鉴定，上述费用待鉴定后再行补充变更"
    else:  # death
        claim_items = f"医疗费、误工费、护理费、住院伙食补助费、营养费、死亡赔偿金（包含被扶养人生活费）、丧葬费、交通费、精神损害抚慰金等暂计人民币{amount_display}"
        retention_clause = ""

    # 诉讼请求第一条：无序号前缀，参考律师模板格式
    p = doc.add_paragraph()
    _set_body_para(p)
    claim_text = f"请求人民法院依法判令被告赔偿原告{claim_items}；"
    if retention_clause:
        claim_text += retention_clause + "；"
    run = p.add_run(claim_text)
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    # 第二条：诉讼费由被告承担（无序号前缀）
    p = doc.add_paragraph()
    _set_body_para(p)
    run = p.add_run("请求人民法院依法判令本案诉讼费由被告承担。")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    # === 事实及理由（5段结构化）===
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_body_para(heading, indent=False)
    run = heading.add_run("事实及理由")
    _set_run_font(run, BODY_FONT, BODY_SIZE)
    run.font.bold = True

    # 第1-5段
    for para_key in ("paragraph_1", "paragraph_2", "paragraph_3", "paragraph_4", "paragraph_5"):
        para_text = context.get(para_key, "")
        if para_text:
            p = doc.add_paragraph(_clean_paragraph_text(para_text))
            _set_body_para(p)
            for run in p.runs:
                _set_run_font(run, BODY_FONT, BODY_SIZE)

    # 结尾固定段：综上所述...
    conclusion = context.get("conclusion_text", "")
    if conclusion:
        p = doc.add_paragraph(_clean_paragraph_text(conclusion))
        _set_body_para(p)
        for run in p.runs:
            _set_run_font(run, BODY_FONT, BODY_SIZE)

    # === 此致 ===
    p = doc.add_paragraph()
    _set_body_para(p)
    run = p.add_run("此致")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    court = context.get("court_name", context.get("受理法院", "")) or ""
    if not court:
        # 从被告名称推断法院：提取"XX县/XX区/XX市"
        defendant_full = context.get("defendant_name", context.get("被告医院全称", "")) or ""
        court_match = re.search(r'([\u4e00-\u9fff]+(?:省|自治区|市|县|区|旗))', defendant_full)
        if court_match:
            region = court_match.group(1)
            # 取最后一个行政区划（最精确的）
            all_regions = re.findall(r'([\u4e00-\u9fff]+(?:县|区|市|旗))', defendant_full)
            if all_regions:
                region = all_regions[-1]
            court = region
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_body_para(p, indent=False)
    run = p.add_run(f"{court}人民法院" if court else "________人民法院")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    # === 具状人 + 日期（同一段落，确保紧邻） ===
    year = context.get("complaint_year", context.get("起诉年份", "")) or ""
    if not year:
        # 如果没有指定年份，使用当前年份
        from datetime import date
        year = str(date.today().year)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_body_para(p, indent=False)
    # 具状人
    run = p.add_run("具状人：")
    _set_run_font(run, BODY_FONT, BODY_SIZE)
    # 换行（不换段）
    run.add_break()
    # 日期：始终使用年份（优先从 context 取，否则用当前年）
    year_cn = _year_to_chinese(year)
    run = p.add_run(f"{year_cn}年   月   日")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# 司法鉴定申请书
# ═══════════════════════════════════════════════════════════════════════════════

def generate_appraisal_application(case_id: str) -> str:
    """生成司法鉴定申请书 Word 文档 → 返回 MinIO key"""
    catalog_data, analysis_result = _get_catalog_and_analysis(case_id)
    doc_bytes = _build_appraisal_docx(catalog_data, analysis_result)
    minio_key = _upload_to_minio(
        case_id, doc_bytes, "司法鉴定申请书.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    logger.info(f"Appraisal application generated: case={case_id} key={minio_key}")
    return minio_key


def generate_appraisal_inline_data(catalog_data: dict, analysis_result: dict) -> bytes | None:
    """纯数据驱动的司法鉴定申请书生成 → 返回 bytes"""
    return _build_appraisal_docx(catalog_data, analysis_result)


def _build_appraisal_docx(catalog_data: dict, analysis_result: dict) -> bytes:
    """构建司法鉴定申请书 DOCX"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    context: dict[str, Any] = dict(analysis_result)
    plaintiffs = _get_plaintiffs(context, catalog_data)

    context.setdefault("applicant_name", plaintiffs[0].get("name", "") if plaintiffs else "")
    context.setdefault("respondent_name", context.get("被告医院全称", context.get("defendant_name", "")))
    context.setdefault("court_name", context.get("受理法院", ""))
    context.setdefault("facts_and_reasons", context.get("事实与理由", "（待补充）"))

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("司法鉴定申请书")
    run.font.size = Pt(22)
    run.font.bold = True

    # 申请人
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(f"申请人：{context.get('applicant_name', '')}")

    # 被申请人
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(f"被申请人：{context.get('respondent_name', '')}")

    # 申请事项
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("申请事项")
    run.font.bold = True

    case_type = context.get("case_type", "injury")
    if case_type == "death":
        p = doc.add_paragraph(
            "1. 被申请人为患者实施的诊疗行为有无过错，是否尽到说明义务、"
            "取得患者或者患者近亲属明确同意的义务；"
        )
        p.paragraph_format.first_line_indent = Cm(0.74)
        p = doc.add_paragraph(
            "2. 若存在过错，被申请人实施的诊疗行为与本案的损害后果之间"
            "是否存在因果关系及原因力大小。"
        )
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p = doc.add_paragraph(
            "1. 被申请人为患者实施的诊疗行为有无过错，是否尽到说明义务、"
            "取得患者或者患者近亲属明确同意的义务；"
        )
        p.paragraph_format.first_line_indent = Cm(0.74)
        p = doc.add_paragraph(
            "2. 若存在过错，患者的伤残等级程度、误工期、护理期、营养期分别是多少；"
        )
        p.paragraph_format.first_line_indent = Cm(0.74)
        p = doc.add_paragraph(
            "3. 被申请人实施的诊疗行为与本案的损害后果之间"
            "是否存在因果关系及原因力大小。"
        )
        p.paragraph_format.first_line_indent = Cm(0.74)

    # 事实与理由
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("事实与理由")
    run.font.bold = True

    facts = context.get("facts_and_reasons", "（待补充）")
    p = doc.add_paragraph(facts)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run("此致")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(6)
    p.add_run(f"{context.get('court_name', '')}人民法院")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"申请人：{context.get('applicant_name', '')}")

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
