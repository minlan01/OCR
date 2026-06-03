"""
民事起诉状 DOCX 生成器
使用 python-docx 生成格式化的 Word 文档
"""
from __future__ import annotations

import io
import re
import threading
from typing import Any, Optional

from loguru import logger
from openai import OpenAI

from config.settings import settings
from services.complaint.template_manager import get_template, get_template_key


_client: Optional[OpenAI] = None
_client_lock = threading.Lock()

# LLM生成失败的占位文本模式
ERROR_PLACEHOLDER_PATTERN = re.compile(r'\[(段落生成失败|生成失败)[^\]]*\]')


def _has_error_placeholder(text: str) -> bool:
    """检查文本是否包含LLM生成失败的占位标记"""
    if not text:
        return False
    return bool(ERROR_PLACEHOLDER_PATTERN.search(text))


def validate_sections_text(sections_text: list[dict]) -> list[str]:
    """验证生成的段落文本中是否有错误占位标记

    Returns:
        list[str]: 失败的段落标题列表，空列表表示全部正常
    """
    failed = []
    for section in sections_text:
        text = section.get("text", "")
        if _has_error_placeholder(text):
            failed.append(section.get("title", section.get("id", "unknown")))
    return failed


def _get_llm_client() -> OpenAI:
    global _client
    if _client is None:
        with _client_lock:
            if _client is None:
                _client = OpenAI(
                    api_key=settings.bailian_api_key_plain,
                    base_url=settings.bailian_text_base_url,
                )
    return _client


def _generate_section_text(
    section: dict[str, Any],
    slot_data: dict[str, dict],
) -> str:
    required_slots = section.get("required_slots", [])
    context_parts = []
    for slot in required_slots:
        data = slot_data.get(slot, {})
        if data:
            import json
            context_parts.append(f"【{slot}信息】\n{json.dumps(data, ensure_ascii=False, indent=2)}")

    if not context_parts:
        return ""

    context = "\n\n".join(context_parts)
    client = _get_llm_client()

    try:
        response = client.chat.completions.create(
            model=settings.bailian_flash_model,
            messages=[
                {
                    "role": "system",
                    "content": [{"type": "text", "text": "你是一位专业的医疗纠纷律师，擅长撰写民事起诉状。请严格按照指令撰写，使用正式的法律文书语言。"}],
                },
                {
                    "role": "user",
                    "content": [{"type": "text", "text": f"{section['prompt']}\n\n参考信息：\n{context}"}],
                },
            ],
            temperature=0.3,
            timeout=settings.bailian_text_timeout,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"Section generation failed for {section['id']}: {e}")
        return f"[生成失败: {section['title']} — {e}]"


def generate_complaint(
    case_type: str,
    is_minor: bool,
    slot_data: dict[str, dict],
) -> bytes:
    template_key = get_template_key(case_type, is_minor)
    template = get_template(template_key)

    sections_text: list[dict[str, str]] = []
    for section in template["sections"]:
        is_optional = section.get("optional", False)
        if is_optional:
            required_slots = section.get("required_slots", [])
            has_data = any(slot_data.get(slot) for slot in required_slots)
            if not has_data:
                continue

        text = _generate_section_text(section, slot_data)
        if text:
            sections_text.append({
                "id": section["id"],
                "title": section["title"],
                "text": text,
            })

    doc_bytes = _build_docx(case_type, is_minor, slot_data, sections_text)
    return doc_bytes


def _build_docx(
    case_type: str,
    is_minor: bool,
    slot_data: dict[str, dict],
    sections: list[dict[str, str]],
) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 格式常量
    BODY_FONT = "仿宋"
    BODY_SIZE = Pt(14)  # 四号
    TITLE_FONT = "黑体"
    TITLE_SIZE = Pt(22)  # 二号
    LEFT_INDENT = Cm(0.05)
    RIGHT_INDENT = Cm(0.2)
    FIRST_INDENT = Cm(0.99)

    def _set_run_font(run, font_name: str, size):
        """设置 run 的字体（含东亚字体）和字号"""
        run.font.name = font_name
        run.font.size = size
        rPr = run._element.get_or_add_rPr()
        rFonts_elem = rPr.find(qn('w:rFonts'))
        if rFonts_elem is None:
            rFonts_elem = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts_elem)
        rFonts_elem.set(qn('w:eastAsia'), font_name)

    def _set_body_para(p, indent=True):
        """设置正文段落格式"""
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = LEFT_INDENT
        p.paragraph_format.right_indent = RIGHT_INDENT
        if indent:
            p.paragraph_format.first_line_indent = FIRST_INDENT

    doc = Document()

    # 默认字体：仿宋四号
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    # 设置默认样式的东亚字体
    style_rPr = style.element.find(qn('w:rPr'))
    if style_rPr is None:
        style_rPr = OxmlElement('w:rPr')
        style.element.append(style_rPr)
    style_rFonts = style_rPr.find(qn('w:rFonts'))
    if style_rFonts is None:
        style_rFonts = OxmlElement('w:rFonts')
        style_rPr.insert(0, style_rFonts)
    style_rFonts.set(qn('w:eastAsia'), BODY_FONT)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # === 标题 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("民事起诉状")
    _set_run_font(run, TITLE_FONT, TITLE_SIZE)
    run.font.bold = True

    plaintiff_data = slot_data.get("plaintiff", {})
    guardian_data = slot_data.get("guardian", {})
    defendant_data = slot_data.get("defendant", {})

    plaintiff_name = plaintiff_data.get("name", "原告")
    defendant_name = defendant_data.get("hospital_name", "被告")

    if is_minor and guardian_data:
        guardian_name = guardian_data.get("guardian_name", "")
        guardian_relation = guardian_data.get("guardian_relation", "")
        plaintiff_line = f"原告：{plaintiff_name}（未成年人），法定代理人/监护人：{guardian_name}（{guardian_relation}）"
    else:
        plaintiff_line = f"原告：{plaintiff_name}"

    p = doc.add_paragraph()
    _set_body_para(p)
    run_label = p.add_run("原告：")
    _set_run_font(run_label, BODY_FONT, BODY_SIZE)
    run_label.font.bold = True
    run_content = p.add_run(plaintiff_line[len("原告："):])
    _set_run_font(run_content, BODY_FONT, BODY_SIZE)

    defendant_line = f"被告：{defendant_name}"
    if defendant_data.get("hospital_address"):
        defendant_line += f"，住所地：{defendant_data['hospital_address']}"
    if defendant_data.get("legal_entity"):
        defendant_line += f"，法定代表人：{defendant_data['legal_entity']}"

    p = doc.add_paragraph()
    _set_body_para(p)
    run_label = p.add_run("被告：")
    _set_run_font(run_label, BODY_FONT, BODY_SIZE)
    run_label.font.bold = True
    run_content = p.add_run(defendant_line[len("被告："):])
    _set_run_font(run_content, BODY_FONT, BODY_SIZE)

    for section in sections:
        # 清理文本中的方括号年份
        text = re.sub(r'\[\d{4}\]', '', section["text"])
        p = doc.add_paragraph(text)
        _set_body_para(p)
        for run in p.runs:
            _set_run_font(run, BODY_FONT, BODY_SIZE)

    fee_data = slot_data.get("fee", {})
    if fee_data and fee_data.get("items"):
        fee_lines = ["赔偿请求："]
        for item in fee_data["items"]:
            name = item.get("name", "")
            amount = item.get("amount", 0)
            if name:
                fee_lines.append(f"  {name}：{amount:,.2f}元" if isinstance(amount, (int, float)) else f"  {name}：{amount}")
        total = fee_data.get("total_amount")
        if total and isinstance(total, (int, float)):
            fee_lines.append(f"  合计：{total:,.2f}元")

        p = doc.add_paragraph("\n".join(fee_lines))
        _set_body_para(p)
        for run in p.runs:
            _set_run_font(run, BODY_FONT, BODY_SIZE)

    p = doc.add_paragraph()
    _set_body_para(p)
    run = p.add_run("此致")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_body_para(p, indent=False)
    run = p.add_run("________________人民法院")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_body_para(p, indent=False)
    run = p.add_run("具状人：________________")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    from datetime import date
    today = date.today()
    year_cn = ''.join({'0':'〇','1':'一','2':'二','3':'三','4':'四',
                       '5':'五','6':'六','7':'七','8':'八','9':'九'}.get(c, c) for c in str(today.year))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_body_para(p, indent=False)
    run = p.add_run(f"{year_cn}年   月    日")
    _set_run_font(run, BODY_FONT, BODY_SIZE)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
