"""
Word 文档导出器

将 ScanStruct 结构化 JSON 结果转换为格式化的 .docx 文件。
支持标题层级、正文段落、表格、列表等元素。
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from loguru import logger

_XML_INVALID_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _clean_text(text: str) -> str:
    return _XML_INVALID_RE.sub("", text).strip()


def _add_run_safe(para, text: str):
    return para.add_run(_clean_text(text))


def export_docx(
    structured: dict[str, Any],
    output_path: Path,
    filename: str = "",
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_styles(doc)

    main_title = filename or "结构化结果"
    if main_title.lower().endswith(".pdf"):
        main_title = main_title[:-4]
    _add_title(doc, main_title)

    source_type = structured.get("source_type", "")
    page_count = structured.get("page_count", 0)
    quality = structured.get("quality", {})
    structure_score = quality.get("structure_score", 0) if quality else 0

    _add_meta_info(doc, source_type, page_count, structure_score)

    sections = structured.get("sections", [])
    if sections:
        _render_sections(doc, sections)

    orphan_paragraphs = structured.get("orphan_paragraphs", [])
    if orphan_paragraphs:
        _add_orphan_paragraphs(doc, orphan_paragraphs)

    lists = structured.get("lists", [])
    if lists:
        _add_lists(doc, lists)

    tables = structured.get("tables", [])
    if tables:
        _add_tables(doc, tables)

    pages = structured.get("pages", [])
    if not sections and not orphan_paragraphs and not lists and pages:
        _add_pages_as_text(doc, pages)

    _add_footer(doc, source_type, filename)

    doc.save(str(output_path))
    logger.info(f"Word document exported: {output_path}")
    return output_path


def export_docx_bytes(
    structured: dict[str, Any],
    filename: str = "",
) -> bytes:
    import atexit
    import os
    import tempfile

    fd, tmp_name = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    tmp_path = Path(tmp_name)

    def _cleanup_atexit():
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass

    atexit.register(_cleanup_atexit)

    try:
        export_docx(structured, tmp_path, filename)
        return tmp_path.read_bytes()
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            logger.warning(f"Failed to clean up temp file: {tmp_path}")
        finally:
            try:
                atexit.unregister(_cleanup_atexit)
            except (TypeError, ValueError):
                pass


def _setup_styles(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)


def _add_title(doc: Document, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = _add_run_safe(para, text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    para.space_after = Pt(12)


def _add_meta_info(doc: Document, source_type: str, page_count: int, structure_score: float):
    rows_data = []
    if source_type:
        label = "文字PDF" if source_type == "text_pdf" else "扫描PDF" if source_type == "scan_pdf" else source_type
        rows_data.append(("来源类型", label))
    if page_count:
        rows_data.append(("总页数", str(page_count)))
    if structure_score:
        rows_data.append(("结构化评分", f"{structure_score:.1%}"))

    if not rows_data:
        return

    table = doc.add_table(rows=len(rows_data), cols=2, style="Table Grid")
    table.autofit = True

    for i, (label, value) in enumerate(rows_data):
        cell_label = table.cell(i, 0)
        cell_label.width = Cm(3)
        p = cell_label.paragraphs[0]
        r = _add_run_safe(p, label)
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.name = "宋体"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        cell_value = table.cell(i, 1)
        p = cell_value.paragraphs[0]
        r = _add_run_safe(p, value)
        r.font.size = Pt(10.5)
        r.font.name = "宋体"
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_paragraph()


def _render_sections(doc: Document, sections: list[dict[str, Any]], depth: int = 0):
    heading_sizes = {1: 16, 2: 15, 3: 14, 4: 12, 5: 12}

    for sec in sections:
        level = sec.get("level", 1)
        title = _clean_text(sec.get("title", ""))
        paragraphs = sec.get("paragraphs", [])
        subsections = sec.get("subsections", [])

        if title:
            para = doc.add_paragraph()
            indent = min(level - 1, 3) * Cm(0.74)
            para.paragraph_format.left_indent = indent
            para.paragraph_format.space_before = Pt(8 if level <= 2 else 4)
            para.paragraph_format.space_after = Pt(4)

            run = _add_run_safe(para, title)
            fs = heading_sizes.get(level, 12)
            run.font.size = Pt(fs)
            run.font.bold = level <= 3
            font_name = "黑体" if level <= 2 else "宋体"
            run.font.name = font_name
            run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        for para_data in paragraphs:
            text = _clean_text(para_data.get("text", ""))
            if not text:
                continue
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(0.74)
            indent = min(level, 3) * Cm(0.74)
            para.paragraph_format.left_indent = indent

            run = _add_run_safe(para, text)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        if subsections:
            _render_sections(doc, subsections, depth=depth + 1)


def _add_orphan_paragraphs(doc: Document, paragraphs: list[dict[str, Any]]):
    if not paragraphs:
        return

    heading_para = doc.add_paragraph()
    run = _add_run_safe(heading_para, "其他段落")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading_para.space_before = Pt(12)

    for para_data in paragraphs:
        text = _clean_text(para_data.get("text", ""))
        if not text:
            continue
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0.74)

        run = _add_run_safe(para, text)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_lists(doc: Document, lists: list[dict[str, Any]]):
    for lst in lists:
        items = lst.get("items", [])

        if not items:
            continue

        for item in items:
            marker = _clean_text(item.get("marker", ""))
            content = _clean_text(item.get("content", ""))
            if not content:
                continue

            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.48)

            if marker:
                run_marker = _add_run_safe(para, f"{marker}. ")
                run_marker.font.size = Pt(12)
                run_marker.font.bold = True
                run_marker.font.name = "宋体"
                run_marker.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

            run = _add_run_safe(para, content)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_tables(doc: Document, tables: list[dict[str, Any]]):
    heading_para = doc.add_paragraph()
    run = _add_run_safe(heading_para, "附表")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    heading_para.space_before = Pt(16)

    for ti, table_data in enumerate(tables):
        caption = _clean_text(table_data.get("caption", f"表 {ti + 1}"))
        if caption:
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = _add_run_safe(cap_para, caption)
            run.font.size = Pt(10.5)
            run.font.bold = True

        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])

        if not headers and not rows:
            continue

        ncols = max(
            len(headers),
            max((len(r) for r in rows), default=0),
            1,
        )

        while len(headers) < ncols:
            headers.append("")

        word_table = doc.add_table(rows=len(rows) + 1, cols=ncols, style="Table Grid")

        for ci, h in enumerate(headers):
            cell = word_table.cell(0, ci)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = _add_run_safe(p, str(h))
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.name = "宋体"
            r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                if ci >= ncols:
                    break
                cell = word_table.cell(ri + 1, ci)
                p = cell.paragraphs[0]
                r = _add_run_safe(p, str(val))
                r.font.size = Pt(10)
                r.font.name = "宋体"
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        doc.add_paragraph()


def _add_pages_as_text(doc: Document, pages: list[dict[str, Any]]):
    for page_data in pages:
        blocks = page_data.get("blocks", [])
        page_num = page_data.get("page", 0)

        if page_num > 1:
            doc.add_page_break()

        for block in blocks:
            text = _clean_text(block.get("text", ""))
            if not text:
                continue
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(0.74)

            run = _add_run_safe(para, text)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_footer(doc: Document, source_type: str, filename: str):
    doc.add_paragraph()
    doc.add_paragraph("—" * 40)

    info_lines = []
    if source_type:
        engine_label = "文字PDF提取" if source_type == "text_pdf" else "OCR识别" if source_type == "scan_pdf" else source_type
        info_lines.append(f"处理方式: {engine_label}")
    if filename:
        info_lines.append(f"源文件: {filename}")

    for line in info_lines:
        para = doc.add_paragraph()
        run = _add_run_safe(para, line)
        run.font.size = Pt(9)
        run.font.name = "宋体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
